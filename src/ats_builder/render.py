from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT


def _setup_document() -> Document:
    doc = Document()
    # Margins
    sect = doc.sections[0]
    sect.top_margin = Inches(0.75)
    sect.bottom_margin = Inches(0.75)
    sect.left_margin = Inches(0.75)
    sect.right_margin = Inches(0.75)
    # Base font
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    try:
        styles["Normal"].paragraph_format.space_after = Pt(0)
        styles["Normal"].paragraph_format.line_spacing = 1.0
    except Exception:
        pass
    # Heading style tweaks
    try:
        styles["Heading 2"].font.name = "Calibri"
        styles["Heading 2"].font.size = Pt(12)
        styles["Heading 2"].paragraph_format.space_before = Pt(6)
        styles["Heading 2"].paragraph_format.space_after = Pt(2)
    except KeyError:
        pass
    return doc


def _add_header(doc: Document, identity: Dict):
    name = identity.get("name", "")
    email = identity.get("email", "")
    phone = identity.get("phone") or ""
    location = identity.get("location") or ""
    links = identity.get("links") or []

    title = doc.add_paragraph()
    run = title.add_run(name)
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    contact = doc.add_paragraph()
    parts = [p for p in [location, email, phone] if p]
    if links:
        parts.extend(links)
    contact.add_run(" | ".join(parts))
    contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact.paragraph_format.space_after = Pt(6)


def _add_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_experience(doc: Document, ctx: Dict):
    _add_heading(doc, "Experience")
    for exp in ctx.get("experiences", []):
        # Role line
        role_line = doc.add_paragraph()
        fmt = role_line.paragraph_format
        fmt.space_after = Pt(0)
        # Right-aligned date via tab stop
        right_edge = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        fmt.tab_stops.add_tab_stop(right_edge, alignment=WD_TAB_ALIGNMENT.RIGHT)
        role_text = f"{exp.get('role','')} | {exp.get('company','')}\t"
        role_run = role_line.add_run(role_text)
        role_run.bold = True
        # Dates on the same line, right-aligned
        if exp.get("start") or exp.get("end"):
            role_line.add_run(f"{exp.get('start','')} – {exp.get('end','')}")
        # Bullets
        for b in exp.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(0)
            bp.add_run(b.get("text", ""))


def _add_education(doc: Document, ctx: Dict):
    if not ctx.get("education"):
        return
    _add_heading(doc, "Education")
    # Deduplicate at render time as well
    seen = set()
    for ed in ctx.get("education", []):
        key = "|".join([
            (ed.get("institution", "") or "").strip().lower(),
            (ed.get("degree", "") or "").strip().lower(),
            str(ed.get("start", "")),
            str(ed.get("end", "")),
        ])
        if key in seen:
            continue
        seen.add(key)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ed.get("institution", ""))
        run.bold = True
        parts = []
        if ed.get("degree"):
            parts.append(ed["degree"])
        if ed.get("start") or ed.get("end"):
            parts.append(f"{ed.get('start','')} – {ed.get('end','')}")
        if parts:
            p.add_run(f" — {' | '.join(parts)}")


def _add_skills(doc: Document, ctx: Dict):
    skills = ctx.get("skills", [])
    if not skills:
        return
    _add_heading(doc, "Skills")
    p = doc.add_paragraph(", ".join(skills))
    p.paragraph_format.space_after = Pt(0)


def render_resume(ctx: Dict, out_path: str) -> str:
    doc = _setup_document()
    _add_header(doc, ctx.get("identity", {}))
    _add_experience(doc, ctx)
    _add_education(doc, ctx)
    _add_skills(doc, ctx)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def render_cover_letter(identity: Dict, letter_ctx: Dict, job_company: Optional[str], out_path: str) -> str:
    doc = _setup_document()
    # Date + Recipient
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph(job_company or "")
    doc.add_paragraph("")
    # Greeting
    doc.add_paragraph(letter_ctx.get("greeting", "Dear Hiring Manager,"))
    # Body
    for para in letter_ctx.get("paragraphs", []):
        doc.add_paragraph(para)
    # Closing
    doc.add_paragraph("")
    doc.add_paragraph(letter_ctx.get("closing", identity.get("name","")))
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def try_export_pdf(docx_path: str) -> Optional[str]:
    p = Path(docx_path)
    pdf_path = p.with_suffix(".pdf")
    # Prefer docx2pdf on Windows
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(str(p), str(pdf_path))
        return str(pdf_path)
    except Exception:
        pass
    # Fallback to LibreOffice headless if available
    try:
        import shutil, subprocess
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            subprocess.run([soffice, "--headless", "--convert-to", "pdf", str(p), "--outdir", str(p.parent)], check=True)
            return str(pdf_path)
    except Exception:
        pass
    return None
