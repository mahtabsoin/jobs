from __future__ import annotations

import re
from dataclasses import asdict
from pathlib import Path
from typing import Dict, List, Tuple

from unidecode import unidecode
from docx import Document as DocxDocument
from pypdf import PdfReader


SECTION_ALIASES = {
    "experience": {"experience", "work experience", "professional experience", "employment"},
    "education": {"education", "academics"},
    "skills": {"skills", "technical skills", "core skills", "tooling"},
    "certifications": {"certifications", "certification", "licenses"},
    "projects": {"projects", "selected projects"},
}

DEGREE_KEYWORDS = [
    "bachelor", "master", "phd", "b.s.", "b.s", "bs", "m.s.", "m.s", "ms", "mba", "ba", "ma",
]


def read_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    chunks: List[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(chunks)


def read_text_from_docx(path: str) -> str:
    doc = DocxDocument(path)
    paras = [p.text for p in doc.paragraphs]
    # Include simple tables content.
    for table in doc.tables:
        for row in table.rows:
            paras.extend([cell.text for cell in row.cells])
    return "\n".join(paras)


def read_text_generic(path: str) -> str:
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        return read_text_from_pdf(str(p))
    if p.suffix.lower() in {".docx"}:
        return read_text_from_docx(str(p))
    return p.read_text(encoding="utf-8", errors="ignore")


def normalize_text(text: str) -> str:
    t = unidecode(text or "")
    t = t.replace("\r", "\n")
    # Collapse excessive blank lines
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t


def detect_sections(lines: List[str]) -> Dict[str, Tuple[int, int]]:
    # Return mapping section_key -> (start_idx, end_idx)
    idxs: Dict[str, Tuple[int, int]] = {}
    headings: List[Tuple[str, int]] = []
    for i, ln in enumerate(lines):
        clean = ln.strip().lower().strip(" :")
        for key, aliases in SECTION_ALIASES.items():
            if clean in aliases or any(clean.startswith(a + ":") for a in aliases):
                headings.append((key, i))
    # Add sentinel end
    headings.sort(key=lambda x: x[1])
    for j, (key, start) in enumerate(headings):
        end = headings[j + 1][1] if j + 1 < len(headings) else len(lines)
        idxs[key] = (start + 1, end)
    return idxs


def parse_bullets(block_lines: List[str]) -> List[str]:
    bullets: List[str] = []
    for ln in block_lines:
        s = ln.strip()
        if not s:
            continue
        if re.match(r"^([\-\*\u2022\u2023\u25E6\u2043\u2219]|\d+\.|\d+\))\s+", s):
            s = re.sub(r"^([\-\*\u2022\u2023\u25E6\u2043\u2219]|\d+\.|\d+\))\s+", "", s)
            bullets.append(s)
    # If no explicit bullet markers, treat each non-empty line as a potential bullet but keep it conservative
    if not bullets and block_lines:
        bullets = [ln.strip() for ln in block_lines if ln.strip()]
    return bullets


def parse_experience(lines: List[str]) -> List[Dict]:
    exps: List[Dict] = []
    # Split by blank lines into chunks
    chunk: List[str] = []
    def flush_chunk():
        nonlocal chunk
        if not chunk:
            return
        # First non-empty line: role/company
        header = next((c for c in chunk if c.strip()), "").strip()
        role = header
        company = ""
        # Try role/company separators
        if " at " in header.lower():
            parts = re.split(r"\s+at\s+", header, flags=re.I)
            if len(parts) == 2:
                role, company = parts[0].strip(), parts[1].strip()
        elif " | " in header:
            parts = header.split(" | ")
            role, company = parts[0].strip(), (parts[1].strip() if len(parts) > 1 else "")
        elif " - " in header:
            parts = header.split(" - ")
            role, company = parts[0].strip(), (parts[1].strip() if len(parts) > 1 else "")

        # Dates: scan lines for year patterns
        text = " \n".join(chunk)
        m = re.findall(r"(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\b\d{4})", text, flags=re.I)
        start = m[0] if m else None
        end = m[-1] if len(m) > 1 else None

        # Bullets: remaining lines
        bullets = parse_bullets(chunk[1:]) if len(chunk) > 1 else []
        if role or company or bullets:
            exps.append({
                "company": company or "",
                "role": role or "",
                "start": start,
                "end": end,
                "bullets": bullets,
            })
        chunk = []

    for ln in lines:
        if ln.strip() == "":
            flush_chunk()
        else:
            chunk.append(ln)
    flush_chunk()
    return exps


def parse_skills(lines: List[str]) -> List[str]:
    text = " ".join(lines)
    parts = re.split(r"[\n,;]\s*", text)
    skills = [p.strip() for p in parts if p.strip()]
    # Deduplicate preserving order
    seen = set()
    out: List[str] = []
    for s in skills:
        k = s.lower()
        if k not in seen:
            seen.add(k)
            out.append(s)
    return out


def parse_education(lines: List[str]) -> List[Dict]:
    out: List[Dict] = []
    chunk: List[str] = []
    def flush():
        nonlocal chunk
        if not chunk:
            return
        text = " ".join(chunk)
        degree = None
        for kw in DEGREE_KEYWORDS:
            if re.search(rf"\b{re.escape(kw)}\b", text, flags=re.I):
                degree = kw
                break
        # Institution: take first line
        institution = chunk[0].strip()
        # Years if present
        m = re.findall(r"\b\d{4}\b", text)
        start = m[0] if m else None
        end = m[1] if len(m) > 1 else None
        out.append({"institution": institution, "degree": degree, "start": start, "end": end})
        chunk = []

    for ln in lines:
        if ln.strip() == "":
            flush()
        else:
            chunk.append(ln)
    flush()
    return out


def parse_resume_text(text: str) -> Dict:
    text = normalize_text(text)
    lines = [ln.rstrip() for ln in text.splitlines()]
    sections = detect_sections(lines)
    experiences = parse_experience(lines[sections["experience"][0]:sections["experience"][1]]) if "experience" in sections else []
    skills = parse_skills(lines[sections["skills"][0]:sections["skills"][1]]) if "skills" in sections else []
    education = parse_education(lines[sections["education"][0]:sections["education"][1]]) if "education" in sections else []
    return {
        "work_experience": experiences,
        "skills": skills,
        "education": education,
    }
